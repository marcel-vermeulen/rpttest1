#! python3
import os
import datetime
#from datetime import datetime, date
import docx

#import pytz
#timezone = pytz.timezone('Africa/Johannesburg')

rptdate = datetime.date.today()
date_time_obj = datetime.datetime.strptime(str(rptdate), '%Y-%m-%d')

print('Date:', date_time_obj.date())
#print('Time:', date_time_obj.time())
#print('Date-time:', date_time_obj)


demodoc = docx.Document('data/demo.docx')
rptdoc = docx.Document()
rptdoc.save('data/test1.docx')

def getText(fdoc):
    fullText = []
    for para in fdoc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

#print(len(demodoc.paragraphs))
#print(demodoc.paragraphs[0].text)
#print(demodoc.paragraphs[1].runs[0].text)

print(getText(demodoc))

print("Preparing Monthly Report")
rptdoc.add_paragraph('Monthly Report', 'Title')
paraObj1 = rptdoc.add_paragraph('Prepared By: ')
paraObj1.add_run(' Marcel Vermeulen')
paraObj2 = rptdoc.add_paragraph('Date Prepared: ')
paraObj2.add_run(str(rptdate))

rptdoc.add_heading('Executive Summary', 1)
execP1txt = rptdoc.add_paragraph('Executive summary text.')

rptdoc.add_heading('Highlights', 3)
execH1txt = rptdoc.add_paragraph('Highlights text.', 'ListBullet')
execH1txt.add_run(' This text is being added to the second paragraph.')

rptdoc.add_heading('Focus Areas', 3)
execH2txt = rptdoc.add_paragraph('Focus text.', 'ListNumber')

rptdoc.add_heading('Risks Table', 3)
execH3txt = rptdoc.add_paragraph('Table summary text.')
table = rptdoc.add_table(rows=1, cols=4)
table.style = 'TableGrid'
# Add headers
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Risk'
hdr_cells[1].text = 'Raised'
hdr_cells[2].text = 'Desc'
hdr_cells[3].text = 'Status'

# get table data -------------
items = (
    ('No water', '2018-06', 'WC running out of water', 'Mitigated'),
    ('No power', '2015-06', 'ZA running out of power', 'Current'),
    ('Corona virus', '2020-02', 'Corona virus pandemic in world', 'Current'),
)

# add a data row for each item
for item in items:
    cells = table.add_row().cells
    cells[0].text = item[0]
    cells[1].text = item[1]
    cells[2].text = item[2]
    cells[3].text = item[3]

rptdoc.save('data/test1.docx')

print("Report Done!")