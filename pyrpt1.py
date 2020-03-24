#! python3
import os
import datetime
#from datetime import datetime, date
import docx
from docx.shared import Pt

import pandas as pd
from esxireleases import get_esxireleases

rptdate = datetime.date.today()
date_time_obj = datetime.datetime.strptime(str(rptdate), '%Y-%m-%d')
print('Date:', date_time_obj.date())

basepath = os.path.abspath(os.path.dirname(__file__))
datapath = os.path.join(basepath,"data")

esxi_rls_json = os.path.join(datapath,"esxiReleases.json")

DOC_TITLE = 'Title'
DOC_SEC_HEAD = 'Heading 1'
DOC_SEC_SUBHEAD = 'Heading 3'
DOC_SEC_SUBSUBHEAD = 'Heading 4'

rptname = 'data/pyrpt1.docx'
rptdoc = docx.Document()
rptdoc.save(rptname)

def style_run(run, is_bold=None, is_italic=None, font_size=11):
    """
    Applies rich text styles to the given Run text object.
    :param Run run: run of text to be styled
    :param bool is_bold: indicates if run should be bold
    :param bool is_italic: indicates if run should be italic
    :param int font_size: indicates font size
    """
    run.font.bold = is_bold
    run.font.italic = is_italic
    run.font.size = Pt(font_size)

# Title Section
def Sec_DocTitle():
	rptdoc.add_paragraph('CPS Monthly Report', DOC_TITLE)
	paraObj1 = rptdoc.add_paragraph('Prepared By: ')
	paraObj1.add_run(' Marcel Vermeulen')
	paraObj2 = rptdoc.add_paragraph('Date Prepared: ')
	paraObj2.add_run(str(rptdate))

# Executive Summary Section
def Sec_ExecSummary():
	rptdoc.add_paragraph('Executive Summary', DOC_SEC_HEAD)
	execP1txt = rptdoc.add_paragraph('Executive summary text.')

	rptdoc.add_paragraph('Highlights', DOC_SEC_SUBHEAD)
	execH1txt = rptdoc.add_paragraph('Highlights text.', 'List Bullet')
	execH1txt.add_run(' This text is being added to the second paragraph.')

	rptdoc.add_paragraph('Focus Areas', DOC_SEC_SUBHEAD)
	execH2txt = rptdoc.add_paragraph('Focus text.', 'List Number')

# Risk Management Section
def Sec_Risks():
	rptdoc.add_paragraph('Risks Management', DOC_SEC_HEAD)
	paragraph = rptdoc.add_paragraph()
	paragraph.add_run('Review of current and emerging risks.').italic = True

	table = rptdoc.add_table(rows=1, cols=4)
	table.style = 'LightShading-Accent1'
	# Add headers
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = 'Risk'
	hdr_cells[1].text = 'Raised'
	hdr_cells[2].text = 'Desc'
	hdr_cells[3].text = 'Status'

	# get table data -------------
	items = (
		('No water', '2018-06', 'WC running out of water', 'Mitigated'),
		('No electricity', '2015-06', 'ZA running out of power', 'Current'),
		('Corona virus', '2020-02', 'Corona virus pandemic in world', 'Current'),
	)

	# add a data row for each item
	for item in items:
		cells = table.add_row().cells
		cells[0].text = item[0]
		cells[1].text = item[1]
		cells[2].text = item[2]
		cells[3].text = item[3]

# Release Management Section
def Sec_ReleaseMgt():
	rptdoc.add_paragraph('Release Management', DOC_SEC_HEAD)
	paragraph = rptdoc.add_paragraph()
	run = paragraph.add_run('Review of current release levels for product life-cycle management.')
	#paraT1txt = rptdoc.add_paragraph('...')
	
	style_run(run=run, is_bold=False, is_italic=True, font_size=11)
	paragraph.add_run(' Added to run...')

	# VMware data
	rptdoc.add_paragraph('VMware', DOC_SEC_SUBHEAD)
	rptdoc.add_paragraph('Highlights text.')
	rptdoc.add_paragraph('This text is being added to the list.', 'List Bullet')
	rptdoc.add_paragraph('This text is also being added to the list.', 'List Bullet')

	esxi_releases = pd.DataFrame()
	esxi_releases = get_esxireleases(esxi_rls_json)
	#print(esxi_releases)

	table = rptdoc.add_table(rows=1, cols=4)
	table.style = 'LightShading-Accent1'
	# Add headers
	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = 'releaseLevel'
	hdr_cells[1].text = 'releaseDate'
	hdr_cells[2].text = 'minorRelease'
	hdr_cells[3].text = 'build'

	items = []
	items = esxi_releases.to_dict('records')

	# add a data row for each item
	for item in items:
		print(item['releaseLevel'])
		cells = table.add_row().cells
		cells[0].text = item['releaseLevel']
		cells[1].text = item['releaseDate']
		cells[2].text = item['minorRelease']
		cells[3].text = item['build']

	# Microsoft data
	rptdoc.add_paragraph('Microsoft', DOC_SEC_SUBHEAD)
	rptdoc.add_paragraph('Highlights text.')
	rptdoc.add_paragraph('Microsoft comment.', 'List Bullet')
	rptdoc.add_paragraph('Another comment.', 'List Bullet')
	#rptdoc.add_paragraph('Another comment.', 'ListNumber')

# Generate Monthly Report
print("Preparing Monthly Report")

Sec_DocTitle()

Sec_ExecSummary()

Sec_Risks()

Sec_ReleaseMgt()

# Save Report to Disk
rptdoc.save('data/pyrpt1.docx')

print("Report Done!")